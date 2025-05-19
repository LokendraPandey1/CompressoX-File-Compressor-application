import zlib
import io
import os
from PIL import Image
import fitz  # PyMuPDF for PDF handling
from docx import Document
import base64
import lzma
import bz2
import gzip
import re
from collections import Counter
import heapq

def compress_text_file(input_path, output_path, algorithm='zlib'):
    """Compress a text file using various algorithms"""
    try:
        # Read the text file
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Get original size
        original_size = os.path.getsize(input_path)
        
        # Choose compression algorithm
        if algorithm == 'zlib':
            compressed_data = zlib.compress(text.encode('utf-8'))
            algo_name = 'Zlib DEFLATE'
            description = 'Uses DEFLATE algorithm with zlib implementation'
        elif algorithm == 'lzma':
            compressed_data = lzma.compress(text.encode('utf-8'))
            algo_name = 'LZMA'
            description = 'Uses Lempel-Ziv-Markov chain algorithm'
        elif algorithm == 'bz2':
            compressed_data = bz2.compress(text.encode('utf-8'))
            algo_name = 'BZIP2'
            description = 'Uses Burrows-Wheeler transform and Huffman coding'
        elif algorithm == 'gzip':
            compressed_data = gzip.compress(text.encode('utf-8'))
            algo_name = 'GZIP'
            description = 'Uses DEFLATE algorithm with gzip format'
        elif algorithm == 'huffman':
            # Custom Huffman implementation
            freq = Counter(text)
            codes = build_huffman_tree(freq)
            compressed_data = huffman_encode(text, codes)
            algo_name = 'Huffman Coding'
            description = 'Uses Huffman coding for character frequency optimization'
        else:
            raise ValueError(f"Unsupported algorithm: {algorithm}")
        
        # Write the compressed data
        with open(output_path, 'wb') as f:
            f.write(compressed_data)
        
        # Get compressed size
        compressed_size = os.path.getsize(output_path)
        
        return {
            'success': True,
            'original_size': original_size,
            'compressed_size': compressed_size,
            'ratio': original_size / compressed_size if compressed_size > 0 else 1,
            'algorithm': algo_name,
            'description': description
        }
        
    except Exception as e:
        print(f"Error compressing text file: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def compress_pdf(input_path, output_path, algorithm='image_optimization', is_lossy=True):
    """Compress a PDF file using various algorithms"""
    try:
        # Validate input file
        if not os.path.exists(input_path):
            return {
                'success': False,
                'error': f'Input file does not exist: {input_path}'
            }
            
        # Open the PDF
        try:
            pdf_document = fitz.open(input_path)
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to open PDF: {str(e)}'
            }
            
        if pdf_document is None:
            return {
                'success': False,
                'error': 'Failed to open PDF document'
            }
            
        # Get original size
        original_size = os.path.getsize(input_path)
        
        if is_lossy:
            if algorithm == 'image_optimization':
                # Optimize images in PDF with quality reduction
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    if page is None:
                        continue
                        
                    image_list = page.get_images()
                    if not image_list:
                        continue
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = pdf_document.extract_image(xref)
                            
                            # Skip if image extraction failed
                            if not base_image or not isinstance(base_image, dict):
                                continue
                                
                            # Get image data safely
                            image_bytes = base_image.get("image")
                            if not image_bytes:
                                continue
                            
                            # Open image with PIL
                            try:
                                image = Image.open(io.BytesIO(image_bytes))
                            except Exception as e:
                                print(f"Error opening image {img_index} on page {page_num}: {str(e)}")
                                continue
                            
                            # Compress image with quality reduction
                            try:
                                buffer = io.BytesIO()
                                image.save(buffer, format='JPEG', quality=60, optimize=True)
                                compressed_image = buffer.getvalue()
                                
                                # Replace image in PDF
                                pdf_document.update_stream(xref, compressed_image)
                            except Exception as e:
                                print(f"Error compressing image {img_index} on page {page_num}: {str(e)}")
                                continue
                                
                        except Exception as e:
                            print(f"Error processing image {img_index} on page {page_num}: {str(e)}")
                            continue
                
                # Save with compression
                try:
                    pdf_document.save(output_path, garbage=4, deflate=True, clean=True)
                except Exception as e:
                    return {
                        'success': False,
                        'error': f'Failed to save compressed PDF: {str(e)}'
                    }
                    
                algo_name = 'PDF Image Quality Reduction'
                description = 'Reduces image quality and optimizes embedded images'
                
            elif algorithm == 'dpi_reduction':
                # Reduce DPI of images
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    if page is None:
                        continue
                        
                    image_list = page.get_images()
                    if not image_list:
                        continue
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = pdf_document.extract_image(xref)
                            
                            # Skip if image extraction failed
                            if not base_image or not isinstance(base_image, dict):
                                continue
                                
                            # Get image data safely
                            image_bytes = base_image.get("image")
                            if not image_bytes:
                                continue
                            
                            # Open image with PIL
                            try:
                                image = Image.open(io.BytesIO(image_bytes))
                            except Exception as e:
                                print(f"Error opening image {img_index} on page {page_num}: {str(e)}")
                                continue
                            
                            # Reduce DPI while maintaining aspect ratio
                            try:
                                width, height = image.size
                                new_width = int(width * 0.7)  # Reduce to 70% of original size
                                new_height = int(height * 0.7)
                                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                                
                                # Save with optimization
                                buffer = io.BytesIO()
                                image.save(buffer, format='JPEG', quality=85, optimize=True)
                                compressed_image = buffer.getvalue()
                                
                                # Replace image in PDF
                                pdf_document.update_stream(xref, compressed_image)
                            except Exception as e:
                                print(f"Error resizing image {img_index} on page {page_num}: {str(e)}")
                                continue
                                
                        except Exception as e:
                            print(f"Error processing image {img_index} on page {page_num}: {str(e)}")
                            continue
                
                try:
                    pdf_document.save(output_path, garbage=4, deflate=True, clean=True)
                except Exception as e:
                    return {
                        'success': False,
                        'error': f'Failed to save compressed PDF: {str(e)}'
                    }
                    
                algo_name = 'PDF DPI Reduction'
                description = 'Reduces image resolution while maintaining readability'
                
            elif algorithm == 'color_reduction':
                # Convert images to grayscale and reduce colors
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    if page is None:
                        continue
                        
                    image_list = page.get_images()
                    if not image_list:
                        continue
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = pdf_document.extract_image(xref)
                            
                            # Skip if image extraction failed
                            if not base_image or not isinstance(base_image, dict):
                                continue
                                
                            # Get image data safely
                            image_bytes = base_image.get("image")
                            if not image_bytes:
                                continue
                            
                            # Open image with PIL
                            try:
                                image = Image.open(io.BytesIO(image_bytes))
                            except Exception as e:
                                print(f"Error opening image {img_index} on page {page_num}: {str(e)}")
                                continue
                            
                            # Convert to grayscale and reduce colors
                            try:
                                image = image.convert('L')  # Convert to grayscale
                                image = image.quantize(colors=64)  # Reduce to 64 colors
                                
                                # Save with optimization
                                buffer = io.BytesIO()
                                image.save(buffer, format='PNG', optimize=True)
                                compressed_image = buffer.getvalue()
                                
                                # Replace image in PDF
                                pdf_document.update_stream(xref, compressed_image)
                            except Exception as e:
                                print(f"Error reducing colors for image {img_index} on page {page_num}: {str(e)}")
                                continue
                                
                        except Exception as e:
                            print(f"Error processing image {img_index} on page {page_num}: {str(e)}")
                            continue
                
                try:
                    pdf_document.save(output_path, garbage=4, deflate=True, clean=True)
                except Exception as e:
                    return {
                        'success': False,
                        'error': f'Failed to save compressed PDF: {str(e)}'
                    }
                    
                algo_name = 'PDF Color Reduction'
                description = 'Converts images to grayscale and reduces color palette'
                
            else:
                raise ValueError(f"Unsupported lossy algorithm: {algorithm}")
                
        else:  # Lossless compression
            if algorithm == 'structure_optimization':
                # Optimize PDF structure without quality loss
                try:
                    pdf_document.save(output_path, 
                                    garbage=4,  # Maximum garbage collection
                                    deflate=True,  # Use DEFLATE compression
                                    clean=True,  # Clean redundant elements
                                    linear=True,  # Optimize for web viewing
                                    pretty=False,  # Minimize whitespace
                                    ascii=False)  # Use binary format
                except Exception as e:
                    return {
                        'success': False,
                        'error': f'Failed to save compressed PDF: {str(e)}'
                    }
                    
                algo_name = 'PDF Structure Optimization'
                description = 'Optimizes PDF structure without quality loss'
                
            elif algorithm == 'object_streams':
                # Use object streams for better compression
                try:
                    pdf_document.save(output_path,
                                    garbage=4,
                                    deflate=True,
                                    clean=True,
                                    linear=True,
                                    pretty=False,
                                    ascii=False,
                                    object_streams=True)  # Enable object streams
                except Exception as e:
                    return {
                        'success': False,
                        'error': f'Failed to save compressed PDF: {str(e)}'
                    }
                    
                algo_name = 'PDF Object Streams'
                description = 'Uses object streams for better compression ratio'
                
            elif algorithm == 'metadata_optimization':
                # Optimize metadata and remove unnecessary elements
                for page in pdf_document:
                    if page is None:
                        continue
                        
                    try:
                        # Remove unnecessary annotations
                        page.delete_annotations()
                        
                        # Clean up text content
                        text = page.get_text()
                        if text and text.strip():
                            # Remove redundant whitespace
                            text = re.sub(r'\s+', ' ', text)
                            page.insert_text((0, 0), text)
                    except Exception as e:
                        print(f"Error processing page: {str(e)}")
                        continue
                
                # Save with maximum optimization
                try:
                    pdf_document.save(output_path,
                                    garbage=4,
                                    deflate=True,
                                    clean=True,
                                    linear=True,
                                    pretty=False,
                                    ascii=False,
                                    object_streams=True)
                except Exception as e:
                    return {
                        'success': False,
                        'error': f'Failed to save compressed PDF: {str(e)}'
                    }
                    
                algo_name = 'PDF Metadata Optimization'
                description = 'Optimizes metadata and removes unnecessary elements'
                
            else:
                raise ValueError(f"Unsupported lossless algorithm: {algorithm}")
        
        # Close the PDF document
        try:
            pdf_document.close()
        except Exception as e:
            print(f"Error closing PDF: {str(e)}")
        
        # Verify output file exists
        if not os.path.exists(output_path):
            return {
                'success': False,
                'error': 'Failed to create output file'
            }
        
        # Get compressed size
        compressed_size = os.path.getsize(output_path)
        
        return {
            'success': True,
            'original_size': original_size,
            'compressed_size': compressed_size,
            'ratio': original_size / compressed_size if compressed_size > 0 else 1,
            'algorithm': algo_name,
            'description': description
        }
        
    except Exception as e:
        print(f"Error compressing PDF: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def compress_docx(input_path, output_path, algorithm='image_optimization'):
    """Compress a DOCX file using various algorithms"""
    try:
        # Open the DOCX
        doc = Document(input_path)
        
        # Get original size
        original_size = os.path.getsize(input_path)
        
        if algorithm == 'image_optimization':
            # Optimize images in document
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//w:drawing'):
                        blip = run._element.xpath('.//a:blip/@r:embed')[0]
                        image_part = doc.part.related_parts[blip]
                        image_bytes = image_part.blob
                        
                        # Open image with PIL
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        # Compress image
                        buffer = io.BytesIO()
                        image.save(buffer, format='JPEG', quality=85, optimize=True)
                        compressed_image = buffer.getvalue()
                        
                        # Replace image in document
                        image_part.blob = compressed_image
            
            algo_name = 'DOCX Image Optimization'
            description = 'Optimizes embedded images using JPEG compression'
            
        elif algorithm == 'text_compression':
            # Compress text content
            for paragraph in doc.paragraphs:
                text = paragraph.text
                compressed_text = zlib.compress(text.encode('utf-8'))
                paragraph.text = compressed_text.decode('utf-8', errors='ignore')
            
            algo_name = 'DOCX Text Compression'
            description = 'Compresses text content using DEFLATE algorithm'
            
        elif algorithm == 'structure_optimization':
            # Optimize document structure
            for paragraph in doc.paragraphs:
                # Remove redundant formatting
                for run in paragraph.runs:
                    if not run.text.strip():
                        run._element.getparent().remove(run._element)
            
            algo_name = 'DOCX Structure Optimization'
            description = 'Optimizes document structure and removes redundant elements'
            
        else:
            raise ValueError(f"Unsupported algorithm: {algorithm}")
        
        # Save the compressed document
        doc.save(output_path)
        
        # Get compressed size
        compressed_size = os.path.getsize(output_path)
        
        return {
            'success': True,
            'original_size': original_size,
            'compressed_size': compressed_size,
            'ratio': original_size / compressed_size if compressed_size > 0 else 1,
            'algorithm': algo_name,
            'description': description
        }
        
    except Exception as e:
        print(f"Error compressing DOCX: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def compress_document(input_path, output_path, file_type, algorithm=None, is_lossy=True):
    """Compress a document file based on its type"""
    try:
        # Validate input parameters
        if not input_path or not output_path:
            return {
                'success': False,
                'error': 'Input and output paths are required'
            }
            
        if not os.path.exists(input_path):
            return {
                'success': False,
                'error': f'Input file does not exist: {input_path}'
            }
            
        # Get file extension
        _, ext = os.path.splitext(input_path)
        ext = ext.lower()
        
        # Validate file type
        if not ext:
            return {
                'success': False,
                'error': 'Could not determine file type from extension'
            }
            
        # If no algorithm specified, choose based on file type and mode
        if algorithm is None:
            if ext == '.txt':
                algorithm = 'zlib' if is_lossy else 'huffman'
            elif ext == '.pdf':
                algorithm = 'image_optimization' if is_lossy else 'structure_optimization'
            elif ext == '.docx':
                algorithm = 'image_optimization' if is_lossy else 'structure_optimization'
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {ext}'
                }
        
        # Compress based on file type
        result = None
        if ext == '.txt':
            result = compress_text_file(input_path, output_path, algorithm)
        elif ext == '.pdf':
            result = compress_pdf(input_path, output_path, algorithm, is_lossy)
        elif ext == '.docx':
            result = compress_docx(input_path, output_path, algorithm)
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {ext}'
            }
            
        # Validate compression result
        if result is None:
            return {
                'success': False,
                'error': 'Compression failed: No result returned'
            }
            
        if not isinstance(result, dict):
            return {
                'success': False,
                'error': 'Compression failed: Invalid result format'
            }
            
        if not result.get('success', False):
            return {
                'success': False,
                'error': result.get('error', 'Unknown compression error')
            }
            
        return result
            
    except Exception as e:
        print(f"Error compressing document: {str(e)}")
        return {
            'success': False,
            'error': f'Compression failed: {str(e)}'
        }

# Helper functions for Huffman coding
def build_huffman_tree(freq):
    """Build a Huffman tree from frequency dictionary"""
    heap = [[weight, [char, ""]] for char, weight in freq.items()]
    heapq.heapify(heap)
    
    while len(heap) > 1:
        lo = heapq.heappop(heap)
        hi = heapq.heappop(heap)
        for pair in lo[1:]:
            pair[1] = '0' + pair[1]
        for pair in hi[1:]:
            pair[1] = '1' + pair[1]
        heapq.heappush(heap, [lo[0] + hi[0]] + lo[1:] + hi[1:])
    
    return dict(sorted(heap[0][1:], key=lambda p: (len(p[1]), p[0])))

def huffman_encode(text, codes):
    """Encode text using Huffman codes"""
    return ''.join(codes[char] for char in text) 